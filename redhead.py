
import os,time
from docx import Document
from docx.shared import *

from mystyle import para_fm,run_fm
from float_picture import *




def add_seal(current_dir):
    #3署名及日期设置
    for file in os.listdir(current_dir):        #对digitfiles添加
        if file.endswith('.docx') and file[:4].isdigit():
            print('正在添加印章：',file)
            doc=Document(file)
            paras=doc.paragraphs
            sign_para_text='未找到署名'
            for para in paras:
                para_text=para.text.replace(' ','')
                if '年' in para_text and '月' in para_text\
                   and '日' in para_text and len(para_text)<12:
                    sign_para = paras[paras.index(para)-1]   #日期上一段
                    print('第{}段有署名：{}'.format(paras.index(para),sign_para.text))
                    sign_para_text=sign_para.text.replace(' ','')
                    try:
                        picture_name = current_dir + '/config/' + sign_para_text + '.png'
                        n=len(para_text)
                        if n == 9:
                            x0=(21-2.6)*28.35-64-7*16/2       #x0是日期的中心坐标
                        elif n ==10:
                            x0=(21-2.6)*28.35-64-7.5*16/2
                        elif n ==11:
                            x0=(21-2.6)*28.35-64-8*16/2
                        else:
                            print("日期有错误")
                        y0 = 28.95*10.5+3.7*28.35
                        x1 = x0-5.7/2*28.35 ; y1 = y0-5.24/2*28.35   #x0是图片中心坐标
                        add_float_picture(para, picture_name , pos_x=Pt(x1), pos_y=Pt(y1-40))  ## 测试插入浮动图片2022.1.9
                        print('印章添加成功。')
                    except:
                        print('Seal不存在：',picture_name)
                    break
                    
            #套红机关名
            para = paras[0].insert_paragraph_before()   #最前段插入发文机关
            para._p.get_or_add_pPr().insert(0,parse_xml('<w:snapToGrid {}  w:val="0"/>'.format(nsdecls('w')))) #取消设置对齐到网格
            run=para.add_run()
            print('署名：',sign_para_text)
            run = para.add_run(sign_para_text+'文件')
            run_fm(run,'方正小标宋简体',72,255,0,0)
            para_fm(para,0,0,1,0,0,0,'C')
            
            para = paras[0].insert_paragraph_before()   #插入空行
            para._p.get_or_add_pPr().insert(0,parse_xml('<w:snapToGrid {}  w:val="0"/>'.format(nsdecls('w')))) #取消设置对齐到网格
            para_fm(para,0,0,28.95,0,0,0,'C')
            
            #插入文号
            print('正在生成文号：',file)
            para = paras[0].insert_paragraph_before()
            para._p.get_or_add_pPr().insert(0,parse_xml('<w:snapToGrid {}  w:val="0"/>'.format(nsdecls('w')))) #取消设置对齐到网格
            run = para.add_run()
            fawenzihao=get_fawenzihao(sign_para_text)
            run.text= fawenzihao[0]
            run_fm(run,'仿宋_GB2312')
            para_fm(para,0,0,28.95,0,0,0,'C')
            print('文号已生成：',run.text)
            
            #插入红色分割线
            para = paras[0].insert_paragraph_before()
            para._p.get_or_add_pPr().insert(0,parse_xml('<w:snapToGrid {}  w:val="0"/>'.format(nsdecls('w')))) #取消设置对齐到网格
            para_fm(para,0,0,28.95,0,0,0,'C')
            
            line_name = current_dir + '/config/' + '红色分割线.png'
            run_fm(run,'仿宋',16)
            add_float_picture(para, line_name , pos_x=Pt(2.8*28.35), pos_y=Pt(10*28.35))
            


            #文档保存docx
            file=str(fawenzihao[1])+file[4:]
            doc.save(file)


            #应用重新打开，调整字体缩放
            ssss=560/(len(sign_para_text)+2)
            import win32com.client as win32
            word = win32.gencache.EnsureDispatch('Word.Application') 
            word.Visible = 0
            doc = word.Documents.Open(current_dir+'/'+file)    #打开新的文档
            doc.Paragraphs(1).Range.Font.Scaling = ssss
            doc.Save()
            doc.Close()
            print('文档已保存：',current_dir+'\\'+file,'\n')

            
def get_fawenzihao(sign_text):
    import json
    with open('./config/发文字号.json','r',encoding='utf-8') as f:
        lines=json.load(f)
    
    try:
        daizi=lines[sign_text]['代字']
    except:
        daizi='代字未找到'
        
    wenhao=1
    year=time.strftime("%Y",time.localtime())
    fawenzihao = daizi +'〔'+ year + '〕'+ str(wenhao) + '号'
    return fawenzihao,wenhao

if __name__ == "__main__":
    current_dir = os.path.abspath('./')
    os.chdir(current_dir)
    add_seal(current_dir)

     