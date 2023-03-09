
from PySide2.QtWidgets import QApplication, QMainWindow, QPushButton,QPlainTextEdit,QMessageBox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os,time

from mystyle import para_fm,run_fm,add_my_styles,clear_styles,my_number_style,set_page
from adds_number import adds_page_number
from redhead import add_seal


def save_docx(doc,doc_name):
    lctime=time.localtime()
    num=time.strftime("%M%S",lctime)                           
    doc.save(num+doc_name)  #保存

def set_headings(doc):
    paras=doc.paragraphs
    title = ['决议','决定','命令','公报','公告','通告','意见',\
                '通知','通报','报告','请示','批复','议案','的函','纪要',\
                '计划','总结','申请','名单','制度','办法','规定','方案','要点']
    heading1=['一、','二、','三、','四、','五、','六、',\
              '七、','八、','九、','十、']
    heading2=['（一）','（二）','（三）','（四）','（五）',\
              '（六）','（七）','（八）','（九）','（十）']
    for para in paras:                          
        if paras.index(para)<3 and len(para.text)<60 and para.text.strip(" ")[-2:] in title:  #标题识别
            para.style = doc.styles['Tit']
            para_fm(para,0,0,28.95,0,0,0,'C')
            for run in para.runs:
                run_fm(run,'方正小标宋简体',22,0,0,0)

        if 1<len(para.text)<30 and '：' in para.text[-1] \
           and len(paras[paras.index(para)-1].text)==0:  #主送识别
            para.style = doc.styles['unindent']
            para_fm(para,0,0,28.95,0,0,0,'J')
            for run in para.runs:
                run_fm(run,'仿宋',16,0,0,0)

        if para.text.strip(" ")[:2] in heading1 and len(para.text)<60 and '：' not in para.text:  #一级标题识别
            para.style = doc.styles['H1']
            for run in para.runs:
                run_fm(run,'黑体',16,0,0,0)

        if para.text.strip(" ")[:3] in heading2 and len(para.text)<60 and '：' not in para.text:   #二级标题识别
            para.style = doc.styles['H2']
            for run in para.runs:
                run_fm(run,'楷体',16,0,0,0)
        
def set_appendix(doc):
    paras=doc.paragraphs
    #2段落遍历完毕后，再次遍历到“附件”段落。设置附件格式，并获得其个数和内容。
    dx_s=[]      #进入循环前，确定卡点变量不被重复覆盖。
    for para in paras:
        dx='附件：'
        if dx in para.text[:5] and len(para.text)>3:
            dx_s.append(paras.index(para))	 #附件段落序列数

    if len(dx_s) !=0:   #如果有附件
        dx_n_strs=[]
        for n in range(len(paras[dx_s[0]+1:])):		#往后每一段，遍历查找附件
            if str(n+1)+'.' in paras[dx_s[0]+n].text[:6]:  #判断n.在不在前6个字符内
                dx_n_str=paras[dx_s[0]+n].text[2:] #获得单个附件字符串
                for p in ['：','.','。']:	  #去掉标点
                    if p in dx_n_str:
                        dx_n_str=dx_n_str.replace(p,'')
                dx_n_strs.append(dx_n_str)	#获得所有附件的字符串

        if len(dx_n_strs)==0:
            dx_n_strs.append(paras[dx_s[0]].text[3:]) #只有一个附件时
            paras[dx_s[0]].style=doc.styles['Apdix']
            para_fm(paras[dx_s[0]],0,0,28.95,80,0,-48,'L')   #段落格式
			
        n=len(dx_n_strs)  #有n个附件
        print('有{}个附件'.format(n))
    
    if len(dx_s) != 0 and n>1:    #如果有1.2.3.
        for i in range(n):
            for j in range(n):		  #遍历所有附件
                if str(i+1)+'.' in paras[dx_s[0]+j].text:
                    paras[dx_s[0]+j].style=doc.styles['Apdix 2']
                    para_f=paras[dx_s[0]+j].paragraph_format   #段落格式赋给para_f
                    para_f.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT  #对齐方式
                    para_f.first_line_indent = Pt(0) 
                    para_f.left_indent = Pt(16*6)   #左缩进（Inches,Cm，Pt）需弥补悬挂负值
                    para_f.first_line_indent = Pt(-16)  #悬挂1个字符
        paras[dx_s[0]].style=doc.styles['Apdix 1']     #重设附件1.的格式
        para_fm(paras[dx_s[0]],0,0,28.95,16*6,0,0,'L')
        para_fm(paras[dx_s[0]],0,0,28.95,16*6,0,-16*4,'L')  #重设附件1.的格式

    #4查找|附件|，标记到[]
    ps=[]
    if len(dx_s) != 0:
        for para in paras[dx_s[0]:len(paras)-1]:
            
            if '附' in para.text and '件' in para.text and len(para.text)<5:   #找到顶格'附件'
                ps.append(paras.index(para))
                print('第{}段有|附件|：{}'.format(ps[-1]+1,para.text.strip('\n')))
                if '：' in para.text[-1]:
                    para.text=para.text[:-1]
                para.style=doc.styles['Blackbody']	#设置顶格附件样式
                para_fm(para,0,0,28.95,0,0,0,'L')
                for run in para.runs:
                    run_fm(run,'黑体',16,0,0,0)	#设置顶格附件格式

    #5取用[]于后续比对
    if len(ps) != 0:
        p=ps[0]
        for para in paras[p:]:   #其下的每一段与附件说明内逐个对比
            for _str in dx_n_strs:
                if len(para.text)>2 and similar(para.text,_str)==1:
                    print('第{}段有附件标题：{}'.format(paras.index(para)+1,para.text))
                    para.style=doc.styles['Tit A']
                    para_fm(para,0,0,28.95,0,0,0,'C')
                    for run in para.runs:
                        run_fm(run,'方正小标宋简体',22,0,0,0)
                                         
def similar(text_a,text_b):
    if type(text_a)==type([]):
        a=set(text_a)
        b=set(text_b)
    elif type(text_a)==type('你好'):
        a=set(list(text_a))
        b=set(list(text_b))
    ab = a & b
    ba = a ^ b
    if len(ab)/(0.1+len(ba))>2:
        return 1
    else:
        return 0

def set_date(doc):
    #3署名及日期设置 
    paras=doc.paragraphs
    for para in paras:
        if '年' in para.text and '月' in para.text\
           and '日' in para.text and len(para.text)<12:
            print('第{}段有日期：{}'.format(paras.index(para)+1,para.text))
            para.style=doc.styles['dater']	#日期格式
            para_fm(para,0,0,28.95,0,16*4,0,'R')	#日期格式
            
            paras[paras.index(para)-1].style = doc.styles['Sign']   #日期上一段的样式
            para_fm_bef=paras[paras.index(para)-1].paragraph_format   #日期上一段的格式
            para_fm_bef.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
            para_fm_bef.first_line_indent = Pt(0)
            b=len(paras[paras.index(para)-1].text)  # b是日期上方署名的字符数
            if len(para.text) in {9,10}:		  #判断日期的字符数
                para_fm_bef.right_indent = Pt((8-0.5*b)*16)
            elif len(para.text)==11:
                para_fm_bef.right_indent = Pt((8.5-0.5*b)*16)

def gennerate_docx():
    doc = Document()
    clear_styles(doc)
    add_my_styles(doc)
    my_number_style(doc)
    set_page(doc)

    text = textEdit.toPlainText()
    for line in text.splitlines():
        doc.add_paragraph(line)
    fileName = text.splitlines()[0]
    set_headings(doc)
    set_appendix(doc)
    set_date(doc)   
    save_docx(doc,f"{fileName}.docx")

    QMessageBox.about(window,'已生成',f'路径：{BASEDIR}')

def openDir():
    os.system("start explorer {BASEDIR}")


if __name__ == '__main__':
    BASEDIR = os.path.dirname(__file__)
    app = QApplication([])

    window = QMainWindow()
    window.resize(500, 400)
    window.move(300, 300)
    window.setWindowTitle('公文生成器')

    textEdit = QPlainTextEdit(window)
    textEdit.setPlaceholderText("请输入正文")
    textEdit.move(10,25)
    textEdit.resize(300,350)

    button1 = QPushButton('普通生成', window)
    button1.move(380,80)
    button1.clicked.connect(gennerate_docx)

    button2 = QPushButton('生成红头文件', window)
    button2.move(380,120)
    button2.clicked.connect(add_seal)

    button3 = QPushButton('打开生成目录', window)
    button3.move(380,160)
    button3.clicked.connect(openDir)

    window.show()
    app.exec_()
